import re
from docx import Document
from typing import Dict
from constants.variables import *
from constants.msg import ErrorType
from .pdf_convertor import Convert2PDF


class DocxTemplatePlaceholder:

    def __init__(self,
                 username: str,
                 template: str,
                 tags: Dict,
):
        self.error = 0
        try:
            self.template_document = Document(template)
            self.file_name = template.split('/')[-1]
            self.replace_tags = self.__prepare_tags(tags)
            self.username = username
        except:
            self.error = ErrorType.missing_doc

    def process(self):
        try:
            self.__process(self.template_document, self.replace_tags)
            path = FILE_FOLDER + self.username + "/" + self.file_name
            self.template_document.save(path)
            return Convert2PDF(path, self.username).DocxToPdf()
        except:
            return ErrorType.missing_doc

    def __process(self, doc, tags):
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                # Todo кароч здесь по кусочкам условия надо ловить
                for regex, replace in tags.items():
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.__process(cell, tags)

    def __prepare_tags(self, tags):
        done_tags = dict()
        for regex, replace in tags.items():
            done_tags[re.compile(fr"<<{regex}>>")] = replace
        return done_tags
