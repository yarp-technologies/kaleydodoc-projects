import shutil
import os
import uuid
from constants.variables import *
from fastapi import UploadFile, File
from docx import Document
import re
import asyncio


def save_file(username: str, input_file_data: UploadFile = File(...)):
    try:
        filename = str(uuid.uuid4())[:8] + '_' + input_file_data.filename
        path = os.path.join(FILE_FOLDER + username, filename)
        with open(f'{path}', "wb") as buffer:
            shutil.copyfileobj(input_file_data.file, buffer)
        return path
    except:
        return None

def get_tags(file_path: str):
    tags = []
    pattern = r"<<(.*?)>>"
    def process(doc):
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                matches = re.findall(pattern, text)
                for match in matches:
                    tags.append(match)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process(cell)

    doc = Document(file_path)
    process(doc)
    return list(set(tags))

def dict_tags(tags):
    tag = {}
    for i in tags:
        tag[i] = ""
    return tag

def transform_user(old: dict, new: dict):
    old.update(new)
    return old

