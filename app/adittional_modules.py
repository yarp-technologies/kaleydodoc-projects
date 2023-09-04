import shutil
import os
import uuid
from constants.variables import *
from fastapi import UploadFile, File
from docx import Document
import re


def save_file(input_file_data: UploadFile = File(...)):
    try:
        filename = str(uuid.uuid4())[:8] + '_' + input_file_data.filename
        path = os.path.join(FILE_FOLDER, filename)
        with open(f'{path}', "wb") as buffer:
            shutil.copyfileobj(input_file_data.file, buffer)
        return path
    except:
        return None

def get_tags(file_path: str):
    tags = []
    pattern = r"<<(.*?)>>"
    def process(doc):
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                matches = re.findall(pattern, text)
                for match in matches:
                    tags.append(match)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process(cell)

    doc = Document(file_path)
    process(doc)
    return tags

def dict_tags(tags):
    tag = {}
    for i in tags:
        tag[i] = ""
    return tag

def prepare_regex(text):
    regex = {}
    if text is None:
        return regex
    reg = text.split("\r\n")
    for tag in reg:
        if tag.find(":") != -1:
            value = tag.replace(" ", "").split(":")
            regex[value[0]] = value[1]
    return regex